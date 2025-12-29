"""Document processing endpoints."""
import uuid
import io
import logging
from typing import Annotated

from fastapi import APIRouter, BackgroundTasks, Body, Depends, File, Form, Header, HTTPException, Path, Query, Request, UploadFile
from fastapi import status as fastapi_status
from pydantic import BaseModel, Field
from sqlalchemy.ext.asyncio import AsyncSession

logger = logging.getLogger(__name__)

from app.agents.router import AgentRouter
from app.agents.types import (
    FlashcardAgentInput,
    FlashcardAgentOutput,
    IngestionAgentInput,
    IngestionAgentOutput,
    KGExtractionAgentInput,
    KGExtractionAgentOutput,
    SummaryAgentInput,
    SummaryAgentOutput,
)
from app.api.dependencies import get_agent_router
from app.infrastructure.database import get_db
from app.schemas.common import AsyncTaskResponse, ErrorResponse
from app.schemas.document import DocumentCreate, DocumentRead
from app.services.agent_run_service import AgentRunService
from app.services.document_service import DocumentService
from app.tasks.agent_tasks import add_agent_task

router = APIRouter()


def get_request_id(x_request_id: Annotated[str | None, Header()] = None) -> str:
    """Extract or generate request ID."""
    import uuid as uuid_lib
    return x_request_id or str(uuid_lib.uuid4())


async def _extract_text_from_file(file: UploadFile, file_content: bytes) -> str:
    """Extract text content from uploaded file.
    
    Args:
        file: Uploaded file object
        file_content: Raw file content bytes
        
    Returns:
        Extracted text content
        
    Raises:
        HTTPException: If file type is unsupported or extraction fails
    """
    file_extension = file.filename.split(".")[-1].lower() if "." in file.filename else ""
    
    if file_extension == "pdf":
        try:
            import pypdf
            pdf_reader = pypdf.PdfReader(io.BytesIO(file_content))
            extracted_text = "\n".join([page.extract_text() for page in pdf_reader.pages])
            if not extracted_text.strip():
                raise HTTPException(
                    status_code=400,
                    detail="PDF file appears to be empty or contains no extractable text",
                )
            return extracted_text
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="PDF extraction requires 'pypdf' library. Install with: pip install pypdf",
            )
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Failed to extract text from PDF: {str(e)}",
            )
    
    elif file_extension in ["doc", "docx"]:
        try:
            from docx import Document as DocxDocument
            docx_file = DocxDocument(io.BytesIO(file_content))
            extracted_text = "\n".join([paragraph.text for paragraph in docx_file.paragraphs])
            if not extracted_text.strip():
                raise HTTPException(
                    status_code=400,
                    detail="DOC/DOCX file appears to be empty or contains no extractable text",
                )
            return extracted_text
        except ImportError:
            logger.warning("python-docx not installed, attempting to read DOC/DOCX as text")
            try:
                return file_content.decode("utf-8", errors="ignore")
            except Exception:
                raise HTTPException(
                    status_code=500,
                    detail="DOC/DOCX extraction requires 'python-docx' library. Install with: pip install python-docx",
                )
        except Exception as e:
            logger.error(f"Error extracting text from DOC/DOCX: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Failed to extract text from DOC/DOCX: {str(e)}",
            )
    
    elif file_extension in ["txt", "md", "text"]:
        try:
            return file_content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return file_content.decode("latin-1")
            except Exception as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Failed to decode text file: {str(e)}",
                )
    else:
        # Try to decode as text for unknown extensions
        try:
            return file_content.decode("utf-8")
        except Exception:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file_extension}. Supported: PDF, DOC, DOCX, TXT, MD",
            )


async def _parse_json_request(http_request: Request) -> tuple[uuid.UUID, uuid.UUID, str | None, str | None, str | None, dict, str]:
    """Parse JSON request body.
    
    Returns:
        Tuple of (workspace_id, user_id, title, doc_type, source_uri, metadata, content)
    """
    body = await http_request.json()
    request = CreateDocumentRequest(**body)
    
    if not request.content:
        raise HTTPException(
            status_code=400,
            detail="'content' field is required in JSON body",
        )
    
    return (
        request.workspace_id,
        request.user_id,
        request.title,
        request.doc_type,
        request.source_url,
        request.metadata or {},
        request.content,
    )


async def _parse_multipart_request(http_request: Request) -> tuple[uuid.UUID, uuid.UUID, str | None, str | None, str | None, dict, str]:
    """Parse multipart/form-data request.
    
    Returns:
        Tuple of (workspace_id, user_id, title, doc_type, source_uri, metadata, extracted_text)
    """
    form = await http_request.form()
    file = form.get("file")
    workspace_id_str = form.get("workspace_id")
    user_id_str = form.get("user_id")
    title = form.get("title")
    
    if not file:
        raise HTTPException(
            status_code=400,
            detail="'file' field is required for file uploads",
        )
    
    if not workspace_id_str or not user_id_str:
        raise HTTPException(
            status_code=400,
            detail="workspace_id and user_id are required when uploading a file (use Form fields)",
        )
    
    try:
        resolved_workspace_id = uuid.UUID(str(workspace_id_str))
        resolved_user_id = uuid.UUID(str(user_id_str))
    except ValueError:
        raise HTTPException(
            status_code=400,
            detail="workspace_id and user_id must be valid UUIDs",
        )
    
    # Read file content
    file_content = await file.read()
    file_extension = file.filename.split(".")[-1].lower() if "." in file.filename else ""
    doc_type = file_extension or "file"
    source_uri = file.filename
    metadata = {"original_filename": file.filename, "file_size": len(file_content)}
    
    # Extract text from file
    extracted_text = await _extract_text_from_file(file, file_content)
    
    if not extracted_text or not extracted_text.strip():
        raise HTTPException(
            status_code=400,
            detail="File appears to be empty or contains no extractable text",
        )
    
    doc_title = str(title) if title else file.filename or "Uploaded Document"
    
    return (
        resolved_workspace_id,
        resolved_user_id,
        doc_title,
        doc_type,
        source_uri,
        metadata,
        extracted_text,
    )


class CreateDocumentRequest(DocumentCreate):
    """Request body for creating a document (extends DocumentCreate with user_id)."""
    
    user_id: uuid.UUID = Field(description="User ID creating the document")


@router.post(
    "/documents",
    response_model=DocumentRead,
    responses={
        201: {"model": DocumentRead},
        400: {"model": ErrorResponse},
        500: {"model": ErrorResponse},
    },
    summary="Create a new document (text or file upload)",
    description="Create a document by either providing text content (JSON) or uploading a file (multipart/form-data). Supports PDF, DOC, TXT, MD files. If preferences.auto_ingest_on_upload=true, ingestion will be triggered automatically.",
    status_code=201,
)
async def create_document(
    http_request: Request,
    background_tasks: BackgroundTasks = None,
    agent_router: Annotated[AgentRouter, Depends(get_agent_router)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
    request_id: Annotated[str, Depends(get_request_id)] = None,
) -> DocumentRead:
    """Create a new document.
    
    Supports two modes:
    1. JSON mode: Send text content in JSON body (Content-Type: application/json)
       - Include workspace_id, user_id, title, content, etc. in JSON
    2. File upload mode: Upload a file (Content-Type: multipart/form-data)
       - Include workspace_id, user_id, title as Form fields
       - Include file as File field
    
    Supported file types:
    - PDF (.pdf) - extracts text using pypdf
    - DOC/DOCX (.doc, .docx) - extracts text using python-docx
    - Text files (.txt, .md, etc.) - read directly
    """
    try:
        # Parse request based on content type
        content_type = http_request.headers.get("content-type", "").lower()
        
        if "application/json" in content_type:
            workspace_id, user_id, title, doc_type, source_uri, metadata, extracted_text = await _parse_json_request(http_request)
        elif "multipart/form-data" in content_type:
            workspace_id, user_id, title, doc_type, source_uri, metadata, extracted_text = await _parse_multipart_request(http_request)
        else:
            raise HTTPException(
                status_code=400,
                detail="Content-Type must be either 'application/json' or 'multipart/form-data'",
            )
        
        # Validate workspace and user exist before creating document
        from app.services.workspace_service import WorkspaceService
        from app.services.user_service import UserService
        
        workspace_service = WorkspaceService(db)
        workspace = await workspace_service.get_workspace(workspace_id)
        if not workspace:
            raise HTTPException(
                status_code=404,
                detail=f"Workspace {workspace_id} not found. Please create a workspace first.",
            )
        
        user_service = UserService(db)
        user = await user_service.get_user_by_id(user_id)
        if not user:
            raise HTTPException(
                status_code=404,
                detail=f"User {user_id} not found. Please sign up first.",
            )
        
        # Create document with text content (service will compute hash and store content)
        document_service = DocumentService(db)
        document = await document_service.create_document(
            workspace_id=workspace_id,
            user_id=user_id,
            title=title,
            source_type=doc_type,
            source_uri=source_uri,
            metadata=metadata,
            raw_text=extracted_text if extracted_text and extracted_text.strip() else None,
        )
        
        # Store text content if not already stored by create_document
        # Note: create_document only computes hash, doesn't store content
        if extracted_text and extracted_text.strip():
            document = await document_service.store_raw_text(document.id, extracted_text)
        
        # Refresh document to ensure all fields are loaded (services already committed)
        # This ensures created_at, updated_at, etc. are populated
        await db.refresh(document)
        
        # Check preferences for auto-ingest
        try:
            from app.services.user_preference_service import UserPreferenceService
            pref_service = UserPreferenceService(db)
            preferences = await pref_service.get_preferences(user_id=user_id)
            
            if preferences.auto_ingest_on_upload and extracted_text:
                # Trigger auto-ingest in background
                input_data = IngestionAgentInput(
                    document_id=document.id,
                    workspace_id=workspace_id,
                    user_id=user_id,
                    raw_text=None,  # Already stored
                )
                agent_run_service = AgentRunService(db)
                input_json = input_data.model_dump(mode='json')  # Convert UUIDs to strings for JSON serialization
                agent_run = await agent_run_service.create_run(
                    workspace_id=workspace_id,
                    user_id=user_id,
                    agent_name="ingestion",
                    input_json=input_json,
                    status="queued",
                )
                document.last_run_id = agent_run.id
                await db.commit()
                await db.refresh(document)
                
                add_agent_task(
                    background_tasks,
                    "ingestion",
                    agent_router.run_ingestion,
                    input_data,
                    agent_run.id,
                    db,
                )
        except Exception as pref_error:
            # If preferences or auto-ingest fails, log but don't fail the document creation
            # Document is already saved by the services
            logger.warning(f"Failed to check preferences or trigger auto-ingest: {str(pref_error)}", exc_info=True)
            # Refresh document one more time to ensure it's up to date
            await db.refresh(document)
        
        # Validate and return document
        try:
            return DocumentRead.model_validate(document)
        except Exception as validation_error:
            logger.error(f"Failed to validate document response [request_id={request_id}]: {str(validation_error)}", exc_info=True)
            # Return basic document info even if validation fails
            raise HTTPException(
                status_code=500,
                detail=f"Error serializing document response [request_id={request_id}]: {str(validation_error)}",
            )
    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error creating document [request_id={request_id}]: {str(e)}",
        )


class IngestDocumentRequest(BaseModel):
    """Request body for document ingestion."""

    workspace_id: uuid.UUID = Field(description="Workspace ID")
    user_id: uuid.UUID = Field(description="User ID")
    raw_text: str | None = Field(default=None, description="Optional raw text to store")


class GenerateFlashcardsRequest(BaseModel):
    """Request body for flashcard generation."""

    workspace_id: uuid.UUID = Field(description="Workspace ID")
    user_id: uuid.UUID = Field(description="User ID")
    mode: str = Field(default="mcq", description="Generation mode: qa or mcq (default: mcq)")


class ExtractKGRequest(BaseModel):
    """Request body for KG extraction."""

    workspace_id: uuid.UUID = Field(description="Workspace ID")
    user_id: uuid.UUID = Field(description="User ID")


# Rate limit placeholder
async def check_rate_limit(
    workspace_id: uuid.UUID, user_id: uuid.UUID, request_id: str
) -> None:
    """Placeholder for rate limiting logic.
    
    TODO: Implement actual rate limiting (e.g., using slowapi or redis)
    """
    # Placeholder: no-op for now
    pass


@router.post(
    "/documents/{document_id}/ingest",
    responses={
        202: {"model": AsyncTaskResponse},
        404: {"model": ErrorResponse},
        409: {"model": ErrorResponse},  # Conflict - ingestion already in progress
        500: {"model": ErrorResponse},
    },
    summary="Ingest and process a document",
    description="Process a document: chunk it and generate embeddings. Always runs asynchronously in background.",
)
async def ingest_document(
    document_id: Annotated[uuid.UUID, Path(description="Document ID to process")],
    request_body: IngestDocumentRequest,
    background_tasks: BackgroundTasks,
    agent_router: Annotated[AgentRouter, Depends(get_agent_router)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
    request_id: Annotated[str, Depends(get_request_id)] = None,
) -> AsyncTaskResponse:
    """Ingest a document using IngestionAgent. Always runs asynchronously."""
    # Rate limit check (placeholder)
    await check_rate_limit(request_body.workspace_id, request_body.user_id, request_id)

    # Idempotency check: prevent duplicate ingestion runs
    document_service = DocumentService(db)
    document = await document_service.get_document(document_id)
    if not document:
        raise HTTPException(status_code=404, detail=f"Document {document_id} not found")
    
    # Check if ingestion is already in progress
    if document.status in ("storing", "chunking", "embedding"):
        raise HTTPException(
            status_code=409,
            detail=f"Ingestion already in progress for document {document_id}. Current status: {document.status}",
        )
    
    # Check for active agent runs for this document
    agent_run_service = AgentRunService(db)
    active_runs = await agent_run_service.get_active_runs(
        workspace_id=request_body.workspace_id,
        agent_name="ingestion",
        document_id=document_id,
    )
    if active_runs:
        raise HTTPException(
            status_code=409,
            detail=f"Ingestion already queued/running for document {document_id}. Run ID: {active_runs[0].id}",
        )

    # Create input
    input_data = IngestionAgentInput(
        document_id=document_id,
        workspace_id=request_body.workspace_id,
        user_id=request_body.user_id,
        raw_text=request_body.raw_text,
    )

    # Always run asynchronously
    try:
        agent_run_service = AgentRunService(db)
        input_json = input_data.model_dump(mode='json')  # Convert UUIDs to strings for JSON serialization
        agent_run = await agent_run_service.create_run(
            workspace_id=request_body.workspace_id,
            user_id=request_body.user_id,
            agent_name="ingestion",
            input_json=input_json,
            status="queued",
        )

        # Add to background tasks
        agent_router = AgentRouter(db)
        add_agent_task(
            background_tasks,
            "ingestion",
            agent_router.run_ingestion,
            input_data,
            agent_run.id,
            db,
        )

        return AsyncTaskResponse(
            run_id=agent_run.id,
            status="queued",
            message="Document ingestion queued. Check agent_runs table for status.",
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error queuing document ingestion [request_id={request_id}]: {str(e)}",
        )


@router.post(
    "/documents/{document_id}/flashcards",
    responses={
        202: {"model": AsyncTaskResponse},
        404: {"model": ErrorResponse},
        500: {"model": ErrorResponse},
    },
    summary="Generate flashcards from a document",
    description="Generate flashcards from a document using FlashcardAgent. Always runs asynchronously in background.",
)
async def generate_flashcards(
    document_id: Annotated[uuid.UUID, Path(description="Source document ID")],
    request_body: GenerateFlashcardsRequest,
    background_tasks: BackgroundTasks,
    agent_router: Annotated[AgentRouter, Depends(get_agent_router)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
    request_id: Annotated[str, Depends(get_request_id)] = None,
) -> AsyncTaskResponse:
    """Generate flashcards from a document using FlashcardAgent. Always runs asynchronously."""
    # Rate limit check (placeholder)
    await check_rate_limit(request_body.workspace_id, request_body.user_id, request_id)

    # Validate mode
    from app.core.constants import FLASHCARD_MODE_TO_CARD_TYPE, FLASHCARD_MODES
    if request_body.mode not in FLASHCARD_MODES:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid mode. Must be one of: {', '.join(FLASHCARD_MODES)}",
        )

    # Map mode to card_type for duplicate check
    card_type = FLASHCARD_MODE_TO_CARD_TYPE[request_body.mode]

    # Check for existing flashcards (duplicate prevention)
    from app.services.flashcard_service import FlashcardService
    flashcard_service = FlashcardService(db)
    existing_flashcards = await flashcard_service.find_existing_flashcards(
        workspace_id=request_body.workspace_id,
        user_id=request_body.user_id,
        document_id=document_id,
        card_type=card_type,
        limit=10,
    )
    
    # Note: We always create a new batch (cards tagged with batch_id)
    # This allows multiple generations while tracking which batch created which cards
    # Future: Add `idempotent=true` query param to return existing cards instead

    # Create input
    input_data = FlashcardAgentInput(
        workspace_id=request_body.workspace_id,
        user_id=request_body.user_id,
        source_document_id=document_id,
        mode=request_body.mode,
    )

    # Always run asynchronously
    try:
        agent_run_service = AgentRunService(db)
        input_json = input_data.model_dump(mode='json')  # Convert UUIDs to strings for JSON serialization
        agent_run = await agent_run_service.create_run(
            workspace_id=request_body.workspace_id,
            user_id=request_body.user_id,
            agent_name="flashcard",
            input_json=input_json,
            status="queued",
        )

        # Add to background tasks
        agent_router = AgentRouter(db)
        add_agent_task(
            background_tasks,
            "flashcard",
            agent_router.run_flashcard,
            input_data,
            agent_run.id,
            db,
        )

        return AsyncTaskResponse(
            run_id=agent_run.id,
            status="queued",
            message="Flashcard generation queued. Check agent_runs table for status.",
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error queuing flashcard generation [request_id={request_id}]: {str(e)}",
        )


@router.post(
    "/documents/{document_id}/kg",
    responses={
        202: {"model": AsyncTaskResponse},
        404: {"model": ErrorResponse},
        500: {"model": ErrorResponse},
    },
    summary="Extract knowledge graph from a document",
    description="Extract concepts and relationships from a document using KGExtractionAgent. Always runs asynchronously in background.",
)
async def extract_kg(
    document_id: Annotated[uuid.UUID, Path(description="Source document ID")],
    request_body: ExtractKGRequest,
    background_tasks: BackgroundTasks,
    agent_router: Annotated[AgentRouter, Depends(get_agent_router)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
    request_id: Annotated[str, Depends(get_request_id)] = None,
) -> AsyncTaskResponse:
    """Extract knowledge graph from a document using KGExtractionAgent. Always runs asynchronously."""
    # Rate limit check (placeholder)
    await check_rate_limit(request_body.workspace_id, request_body.user_id, request_id)

    # Create input
    input_data = KGExtractionAgentInput(
        workspace_id=request_body.workspace_id,
        user_id=request_body.user_id,
        source_document_id=document_id,
    )

    # Always run asynchronously
    try:
        agent_run_service = AgentRunService(db)
        input_json = input_data.model_dump(mode='json')  # Convert UUIDs to strings for JSON serialization
        agent_run = await agent_run_service.create_run(
            workspace_id=request_body.workspace_id,
            user_id=request_body.user_id,
            agent_name="kg_extraction",
            input_json=input_json,
            status="queued",
        )

        # Add to background tasks
        agent_router = AgentRouter(db)
        add_agent_task(
            background_tasks,
            "kg_extraction",
            agent_router.run_kg_extraction,
            input_data,
            agent_run.id,
            db,
        )

        return AsyncTaskResponse(
            run_id=agent_run.id,
            status="queued",
            message="KG extraction queued. Check agent_runs table for status.",
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error queuing KG extraction [request_id={request_id}]: {str(e)}",
        )


# Workspace-scoped document routes (contract requirement)
@router.post(
    "/workspaces/{workspace_id}/documents",
    response_model=DocumentRead,
    status_code=201,
    responses={400: {"model": ErrorResponse}, 500: {"model": ErrorResponse}},
    summary="Create a document in a workspace (text or file upload)",
    description="Create a document by either providing text content (JSON) or uploading a file (multipart/form-data). Supports PDF, DOC, TXT, MD files. If preferences.auto_ingest_on_upload=true, ingestion will be triggered automatically.",
)
async def create_workspace_document(
    workspace_id: Annotated[uuid.UUID, Path(description="Workspace ID")],
    # For JSON requests (text content)
    request: Annotated[CreateDocumentRequest | None, Body()] = None,
    # For file uploads (multipart/form-data)
    file: Annotated[UploadFile | None, File(description="File to upload (PDF, DOC, TXT, MD, etc.)")] = None,
    user_id: Annotated[uuid.UUID | None, Form(description="User ID (required for file uploads)")] = None,
    title: Annotated[str | None, Form(description="Document title (optional for file uploads)")] = None,
    background_tasks: BackgroundTasks = None,
    agent_router: Annotated[AgentRouter, Depends(get_agent_router)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
    request_id: Annotated[str, Depends(get_request_id)] = None,
) -> DocumentRead:
    """Create a document in a workspace.
    
    Supports two modes:
    1. JSON mode: Send text content in JSON body (Content-Type: application/json)
    2. File upload mode: Upload a file (Content-Type: multipart/form-data)
    
    Supported file types:
    - PDF (.pdf) - extracts text using pypdf
    - DOC/DOCX (.doc, .docx) - extracts text (requires python-docx)
    - Text files (.txt, .md, etc.) - read directly
    """
    try:
        extracted_text = None
        doc_title = None
        doc_type = None
        source_uri = None
        metadata = {}
        resolved_user_id = None
        
        # Determine mode: file upload or JSON text
        if file:
            # File upload mode
            if not user_id:
                raise HTTPException(
                    status_code=400,
                    detail="user_id is required when uploading a file (use Form field)",
                )
            
            resolved_user_id = user_id
            doc_title = title or file.filename or "Uploaded Document"
            
            # Read file content
            file_content = await file.read()
            file_extension = file.filename.split(".")[-1].lower() if "." in file.filename else ""
            doc_type = file_extension or "file"
            source_uri = file.filename
            metadata = {"original_filename": file.filename, "file_size": len(file_content)}
            
            # Extract text based on file type
            if file_extension == "pdf":
                try:
                    import pypdf
                    pdf_reader = pypdf.PdfReader(io.BytesIO(file_content))
                    extracted_text = "\n".join([page.extract_text() for page in pdf_reader.pages])
                    if not extracted_text.strip():
                        raise HTTPException(
                            status_code=400,
                            detail="PDF file appears to be empty or contains no extractable text",
                        )
                except ImportError:
                    raise HTTPException(
                        status_code=500,
                        detail="PDF extraction requires 'pypdf' library. Install with: pip install pypdf",
                    )
                except Exception as e:
                    logger.error(f"Error extracting text from PDF: {str(e)}")
                    raise HTTPException(
                        status_code=400,
                        detail=f"Failed to extract text from PDF: {str(e)}",
                    )
            elif file_extension in ["doc", "docx"]:
                try:
                    from docx import Document as DocxDocument
                    docx_file = DocxDocument(io.BytesIO(file_content))
                    extracted_text = "\n".join([paragraph.text for paragraph in docx_file.paragraphs])
                    if not extracted_text.strip():
                        raise HTTPException(
                            status_code=400,
                            detail="DOC/DOCX file appears to be empty or contains no extractable text",
                        )
                except ImportError:
                    # Fallback: try to read as text if python-docx not available
                    logger.warning("python-docx not installed, attempting to read DOC/DOCX as text")
                    try:
                        extracted_text = file_content.decode("utf-8", errors="ignore")
                    except Exception:
                        raise HTTPException(
                            status_code=500,
                            detail="DOC/DOCX extraction requires 'python-docx' library. Install with: pip install python-docx",
                        )
                except Exception as e:
                    logger.error(f"Error extracting text from DOC/DOCX: {str(e)}")
                    raise HTTPException(
                        status_code=400,
                        detail=f"Failed to extract text from DOC/DOCX: {str(e)}",
                    )
            elif file_extension in ["txt", "md", "text"]:
                # Read as text
                try:
                    extracted_text = file_content.decode("utf-8")
                except UnicodeDecodeError:
                    try:
                        extracted_text = file_content.decode("latin-1")
                    except Exception as e:
                        raise HTTPException(
                            status_code=400,
                            detail=f"Failed to decode text file: {str(e)}",
                        )
            else:
                # Try to decode as text for unknown extensions
                try:
                    extracted_text = file_content.decode("utf-8")
                except Exception:
                    raise HTTPException(
                        status_code=400,
                        detail=f"Unsupported file type: {file_extension}. Supported: PDF, DOC, DOCX, TXT, MD",
                    )
            
            if not extracted_text or not extracted_text.strip():
                raise HTTPException(
                    status_code=400,
                    detail="File appears to be empty or contains no extractable text",
                )
        
        elif request:
            # JSON text mode
            if request.workspace_id != workspace_id:
                raise HTTPException(status_code=400, detail="Workspace ID mismatch")
            
            resolved_user_id = request.user_id
            doc_title = request.title
            doc_type = request.doc_type
            source_uri = request.source_url
            metadata = request.metadata or {}
            extracted_text = request.content
        
        else:
            raise HTTPException(
                status_code=400,
                detail="Either provide 'file' (multipart/form-data) or JSON body with 'content' field",
            )
        
        # Create document
        document_service = DocumentService(db)
        document = await document_service.create_document(
            workspace_id=workspace_id,
            user_id=resolved_user_id,
            title=doc_title,
            source_type=doc_type,
            source_uri=source_uri,
            metadata=metadata,
        )
        
        # Store text content if available
        if extracted_text:
            document = await document_service.store_raw_text(document.id, extracted_text)
        
        # Check preferences for auto-ingest
        from app.services.user_preference_service import UserPreferenceService
        pref_service = UserPreferenceService(db)
        preferences = await pref_service.get_preferences(user_id=resolved_user_id)
        
        if preferences.auto_ingest_on_upload and extracted_text:
            # Trigger auto-ingest in background
            input_data = IngestionAgentInput(
                document_id=document.id,
                workspace_id=workspace_id,
                user_id=resolved_user_id,
                raw_text=None,  # Already stored
            )
            agent_run_service = AgentRunService(db)
            input_json = input_data.model_dump(mode='json')  # Convert UUIDs to strings for JSON serialization
            agent_run = await agent_run_service.create_run(
                workspace_id=workspace_id,
                user_id=resolved_user_id,
                agent_name="ingestion",
                input_json=input_json,
                status="queued",
            )
            document.last_run_id = agent_run.id
            await db.commit()
            
            add_agent_task(
                background_tasks,
                "ingestion",
                agent_router.run_ingestion,
                input_data,
                agent_run.id,
                db,
            )
        
        return DocumentRead.model_validate(document)
    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error creating document [request_id={request_id}]: {str(e)}",
        )


@router.get(
    "/workspaces/{workspace_id}/documents",
    response_model=list[DocumentRead],
    responses={500: {"model": ErrorResponse}},
    summary="List documents in a workspace",
)
async def list_workspace_documents(
    workspace_id: Annotated[uuid.UUID, Path(description="Workspace ID")],
    db: Annotated[AsyncSession, Depends(get_db)] = None,
) -> list[DocumentRead]:
    """List all documents in a workspace."""
    try:
        document_service = DocumentService(db)
        documents = await document_service.list_documents(workspace_id=workspace_id)
        return [DocumentRead.model_validate(d) for d in documents]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error listing documents: {str(e)}")


@router.get(
    "/documents/{document_id}",
    response_model=DocumentRead,
    responses={404: {"model": ErrorResponse}, 500: {"model": ErrorResponse}},
    summary="Get a document",
)
async def get_document(
    document_id: Annotated[uuid.UUID, Path(description="Document ID")],
    db: Annotated[AsyncSession, Depends(get_db)] = None,
) -> DocumentRead:
    """Get a document by ID (includes status, summary_text, last_run_id)."""
    try:
        document_service = DocumentService(db)
        document = await document_service.get_document(document_id)
        if not document:
            raise HTTPException(status_code=404, detail=f"Document {document_id} not found")
        return DocumentRead.model_validate(document)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting document: {str(e)}")


@router.patch(
    "/documents/{document_id}",
    response_model=DocumentRead,
    responses={404: {"model": ErrorResponse}, 500: {"model": ErrorResponse}},
    summary="Update a document",
)
async def update_document(
    document_id: Annotated[uuid.UUID, Path(description="Document ID")],
    request: DocumentCreate,  # Reuse for partial update
    db: Annotated[AsyncSession, Depends(get_db)] = None,
) -> DocumentRead:
    """Update a document."""
    try:
        document_service = DocumentService(db)
        document = await document_service.update_document(
            document_id=document_id,
            title=request.title,
            doc_type=request.doc_type,
            source_url=request.source_url,
            language=request.language,
            metadata=request.metadata,
        )
        return DocumentRead.model_validate(document)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error updating document: {str(e)}")


@router.delete(
    "/documents/{document_id}",
    status_code=204,
    responses={404: {"model": ErrorResponse}, 500: {"model": ErrorResponse}},
    summary="Delete a document",
)
async def delete_document(
    document_id: Annotated[uuid.UUID, Path(description="Document ID")],
    db: Annotated[AsyncSession, Depends(get_db)] = None,
) -> None:
    """Delete a document (cascade deletes chunks, embeddings, etc.)."""
    try:
        document_service = DocumentService(db)
        await document_service.delete_document(document_id)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}")


@router.get(
    "/documents/{document_id}/status",
    response_model=DocumentRead,
    responses={404: {"model": ErrorResponse}, 500: {"model": ErrorResponse}},
    summary="Get document status",
)
async def get_document_status(
    document_id: Annotated[uuid.UUID, Path(description="Document ID")],
    db: Annotated[AsyncSession, Depends(get_db)] = None,
) -> DocumentRead:
    """Get document status (alias to document GET)."""
    return await get_document(document_id, db)


@router.get(
    "/documents/{document_id}/summary",
    responses={404: {"model": ErrorResponse}, 500: {"model": ErrorResponse}},
    summary="Get document summary",
)
async def get_document_summary(
    document_id: Annotated[uuid.UUID, Path(description="Document ID")],
    db: Annotated[AsyncSession, Depends(get_db)] = None,
) -> dict:
    """Get document summary."""
    try:
        document_service = DocumentService(db)
        document = await document_service.get_document(document_id)
        if not document:
            raise HTTPException(status_code=404, detail=f"Document {document_id} not found")
        return {"summary": document.summary_text, "document_id": str(document_id)}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting summary: {str(e)}")


@router.post(
    "/documents/{document_id}/summary",
    responses={
        202: {"model": AsyncTaskResponse},
        404: {"model": ErrorResponse},
        500: {"model": ErrorResponse},
    },
    summary="Regenerate document summary",
    description="Generate or regenerate document summary using SummaryAgent. Always runs asynchronously in background.",
)
async def regenerate_document_summary(
    document_id: Annotated[uuid.UUID, Path(description="Document ID")],
    background_tasks: BackgroundTasks,
    max_bullets: Annotated[
        int,
        Query(
            description="Maximum number of bullet points",
            ge=1,
            le=20
        )
    ] = 7,  # TODO: Use DEFAULT_SUMMARY_MAX_BULLETS from constants
    agent_router: Annotated[AgentRouter, Depends(get_agent_router)] = None,
    db: Annotated[AsyncSession, Depends(get_db)] = None,
    request_id: Annotated[str, Depends(get_request_id)] = None,
) -> AsyncTaskResponse:
    """Regenerate document summary using SummaryAgent. Always runs asynchronously."""
    try:
        # Get document to extract workspace_id and user_id
        document_service = DocumentService(db)
        document = await document_service.get_document(document_id)
        if not document:
            raise HTTPException(status_code=404, detail=f"Document {document_id} not found")
        
        # Create input
        input_data = SummaryAgentInput(
            document_id=document_id,
            workspace_id=document.workspace_id,
            user_id=document.user_id,  # Use document user_id
            max_bullets=max_bullets,
        )

        # Always run asynchronously
        try:
            agent_run_service = AgentRunService(db)
            input_json = input_data.model_dump(mode='json')  # Convert UUIDs to strings for JSON serialization
            agent_run = await agent_run_service.create_run(
                workspace_id=document.workspace_id,
                user_id=document.user_id,
                agent_name="summary",
                input_json=input_json,
                status="queued",
            )

            # Add to background tasks
            agent_router = AgentRouter(db)
            add_agent_task(
                background_tasks,
                "summary",
                agent_router.run_summary,
                input_data,
                agent_run.id,
                db,
            )

            return AsyncTaskResponse(
                run_id=agent_run.id,
                status="queued",
                message="Summary generation queued. Check agent_runs table for status.",
            )
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error queuing summary generation [request_id={request_id}]: {str(e)}",
            )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error regenerating summary: {str(e)}")


@router.post(
    "/documents/{document_id}/reindex",
    responses={
        200: {"description": "Reindex successful"},
        404: {"model": ErrorResponse},
        500: {"model": ErrorResponse},
    },
    summary="Reindex document embeddings",
    description="Delete old embeddings and regenerate with current embedding model. Useful after changing embedding model configuration.",
)
async def reindex_document(
    document_id: Annotated[uuid.UUID, Path(description="Document ID to reindex")],
    embedding_model: Annotated[str, Query(description="Embedding model to use")] = "default",
    db: Annotated[AsyncSession, Depends(get_db)] = None,
) -> dict:
    """Reindex document embeddings.
    
    This endpoint:
    - Deletes old embeddings from Qdrant (by document_id filter)
    - Deletes old Embedding records from DB
    - Regenerates embeddings with the specified model
    - Upserts new vectors to Qdrant
    
    Chunk IDs remain stable, ensuring chat and other features continue to work.
    """
    try:
        from app.services.embedding_service import EmbeddingService
        from app.infrastructure.qdrant import QdrantClientWrapper
        
        document_service = DocumentService(db)
        document = await document_service.get_document(document_id)
        if not document:
            raise HTTPException(status_code=404, detail=f"Document {document_id} not found")
        
        # Reindex embeddings
        qdrant_client = QdrantClientWrapper()
        embedding_service = EmbeddingService(db, qdrant_client=qdrant_client)
        new_embeddings = await embedding_service.reindex_document(
            document_id=document_id,
            embedding_model=embedding_model,
        )
        
        return {
            "document_id": str(document_id),
            "embeddings_created": len(new_embeddings),
            "embedding_model": embedding_model,
            "message": "Document reindexed successfully. Old vectors replaced with new ones.",
        }
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Error reindexing document: {str(e)}"
        )

